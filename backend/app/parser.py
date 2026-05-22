import os
import pdfplumber
from pypdf import PdfReader
from typing import Tuple

def extract_text_from_txt(file_path: str) -> str:
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read().strip()

def extract_text_from_pdf(file_path: str) -> Tuple[str, bool]:
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            pages_text = []
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text(layout=True)
                if page_text:
                    pages_text.append(page_text)
            text = "\n--- Page Separator ---\n".join(pages_text)
    except Exception as e:
        print(f"pdfplumber failed: {e}. Falling back to pypdf.")
        text = ""

    if not text.strip():
        try:
            reader = PdfReader(file_path)
            pages_text = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    pages_text.append(page_text)
            text = "\n--- Page Separator ---\n".join(pages_text)
        except Exception as e:
            print(f"pypdf fallback failed: {e}")
            text = ""

    is_scanned = len(text.strip()) < 50
    return text.strip(), is_scanned

def extract_text_from_docx(file_path: str) -> str:
    try:
        import docx
    except ImportError:
        raise ImportError("python-docx package is not installed. Add it to requirements.txt.")
        
    doc = docx.Document(file_path)
    content = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text)
            
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                content.append(" | ".join(row_cells))
                
    return "\n".join(content).strip()

def extract_content(file_path: str) -> Tuple[str, bool]:
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.txt':
        return extract_text_from_txt(file_path), False
    elif ext == '.docx':
        return extract_text_from_docx(file_path), False
    elif ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.png', '.jpg', '.jpeg']:
        return "", True
    else:
        raise ValueError(f"Unsupported file extension: {ext}")
